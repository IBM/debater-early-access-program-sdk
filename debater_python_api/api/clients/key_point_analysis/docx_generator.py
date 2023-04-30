import logging
import random

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor
from docx import Document
from debater_python_api.api.clients.key_point_analysis.utils import create_dict_to_list, \
    trunc_float, read_dicts_from_df, get_unique_sent_id

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

def add_link(paragraph, link_to, text, tool_tip=None, set_color=False):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to,)

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip,)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.bold = False
    r.italic = False
    if set_color:
        r.font.color.rgb = RGBColor(0x40, 0x5E, 0x8D)

def set_heading(heading):
    paragraph_format = heading.paragraph_format
    paragraph_format.page_break_before = False
    paragraph_format.keep_with_next = False
    paragraph_format.keep_together = False


def get_unique_matches_subtree(node_id, id_to_node, id_to_kids, id_to_n_unique_matches_subtree, kp_to_dicts, by_sentence = True):
    kp = id_to_node[node_id]['data']['kp']
    if by_sentence:
        kp_unique_sentences = set([get_unique_sent_id(d) for d in kp_to_dicts[kp]])
    else: # by comments
        kp_unique_sentences = set([d['comment_id'] for d in kp_to_dicts[kp]])
    if node_id in id_to_kids:
        for kid in id_to_kids[node_id]:
            kp_unique_sentences = kp_unique_sentences.union(get_unique_matches_subtree(kid, id_to_node, id_to_kids, id_to_n_unique_matches_subtree, kp_to_dicts, by_sentence))
    id_to_n_unique_matches_subtree[node_id] = len(kp_unique_sentences)
    return kp_unique_sentences

def add_data_stats(dicts, nodes_ids, document, stance, min_n_matches):
    dicts_not_none = list(filter(lambda r: r["kp"] != "none", dicts))
    n_sents = len(set([get_unique_sent_id(d) for d in dicts]))
    rate_matched_sents = 100 *len(set([get_unique_sent_id(d) for d in dicts_not_none])) / n_sents
    n_comments = len(set([d["comment_id"] for d in dicts]))
    rate_matched_comments = 100 *len(set([d["comment_id"] for d in dicts_not_none])) / n_comments
    n_kps = len(nodes_ids)
    if stance == "pos":
        stance_str = " with positive sentiment"
    elif stance == "neg":
        stance_str = " with negative sentiment"
    else:
        stance_str = ""

    heading = document.add_heading('Data Statistics', 1)
    set_heading(heading)
    p = document.add_paragraph()
    run = p.add_run(
         f'Analyzed {n_comments} comments ({n_sents} sentences){stance_str}.\n'
        f'Identified {n_kps} key points with at least {min_n_matches or 1} matching sentences.\n'
        f'{int(rate_matched_comments)}% of the comments (and {int(rate_matched_sents)}%'
        f' of the sentences) were matched to at least one key point.'
        )

    run.font.size = Pt(12)


def sample_list_keep_order(list_to_sample, n_to_sample):
    return [list_to_sample[i] for i in sorted(random.sample(range(len(list_to_sample)), n_to_sample))]


def save_hierarchical_graph_data_to_docx(full_result_df, graph_data, result_filename, n_matches=None, sort_by_subtree=True, include_match_score=False, min_n_matches=5):
    def get_hierarchical_bullets_aux(document, id_to_kids, id_to_node, id, tab, id_to_paragraph, id_to_n_matches_subtree, sort_by_subtree=True, ids_order=[]):
        bullet = '\u25E6' if tab % 2 == 1 else '\u2022'
        msg = f'{("   " * tab)} {bullet} '

        heading = document.add_heading(msg, 9 if (tab + 1) > 9 else (tab+1))
        set_heading(heading)

        id_to_paragraph[id] = heading
        ids_order.append(id)

        if id in id_to_kids:
            kids = id_to_kids[id]
            if sort_by_subtree:
                kids = sorted(kids, key=lambda n: int(id_to_n_matches_subtree[n]), reverse=True)
            else:
                kids = sorted(kids, key=lambda n: int(id_to_node[n]['data']['n_matches']), reverse=True)
            for k in kids:
                get_hierarchical_bullets_aux(document, id_to_kids, id_to_node, k, tab + 1, id_to_paragraph, id_to_n_matches_subtree, sort_by_subtree, ids_order)


    def get_hierarchical_bullets(document, roots, id_to_kids, id_to_node, id_to_paragraph, id_to_n_matches_subtree, sort_by_subtree=True, ids_order=[]):
        tab = 0
        if sort_by_subtree:
            roots = sorted(roots, key=lambda n: int(id_to_n_matches_subtree[n]), reverse=True)
        else:
            roots = sorted(roots, key=lambda n: int(id_to_node[n]['data']['n_matches']), reverse=True)
        for root in roots:
            get_hierarchical_bullets_aux(document, id_to_kids, id_to_node, root, tab, id_to_paragraph, id_to_n_matches_subtree, sort_by_subtree=True, ids_order=ids_order)

    nodes = [d for d in graph_data if d['type'] == 'node']
    edges = [d for d in graph_data if d['type'] == 'edge']

    if min_n_matches is not None:
        nodes = [n for n in nodes if int(n['data']['n_matches'])>=min_n_matches]

    nodes_ids = [n['data']['id'] for n in nodes]
    edges = [e for e in edges if (e['data']['target'] in nodes_ids and e['data']['source'] in nodes_ids)]

    id_to_kids = create_dict_to_list([(e['data']['target'], e['data']['source']) for e in edges])
    id_to_node = {d['data']['id']: d for d in nodes}

    all_kids = set()
    for id, kids in id_to_kids.items():
        all_kids = all_kids.union(kids)

    root_ids = set(nodes_ids).difference(all_kids)

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'

    dicts, _ = read_dicts_from_df(full_result_df)
    kp_to_dicts = create_dict_to_list([(d['kp'], d) for d in dicts])

    id_to_n_matches_subtree = {}
    for root in root_ids:
        get_unique_matches_subtree(root, id_to_node, id_to_kids, id_to_n_matches_subtree, kp_to_dicts)

    stances = set([d['stance'] for d in dicts if 'stance' in d])
    stances = stances.union(set([d['selected_stance'] for d in dicts if 'selected_stance' in d]))

    stance = None
    if len(stances) == 1 and 'pos' in stances:
        stance = 'pos'
    elif (len(stances) == 1 and ('neg' in stances or 'sug' in stances)) \
            or (len(stances) == 2 and 'neg' in stances and 'sug' in stances):
        stance = 'neg'

    heading = document.add_heading('Key Point Analysis Results', 1)
    set_heading(heading)

    if stance == 'pos':
        p = document.add_paragraph('Positive Key Points')
        p.style = document.styles['Subtitle']
        set_heading(p)
        insertHR(p)
    elif stance == 'neg':
        p = document.add_paragraph('Negative Key Points')
        p.style = document.styles['Subtitle']
        set_heading(p)
        insertHR(p)
    else:
        insertHR(heading)

    add_data_stats(dicts, nodes_ids, document, stance, min_n_matches)

    heading = document.add_heading('Key Point Hierarchy', 1)
    set_heading(heading)

    p = document.add_paragraph()
    run = p.add_run('Click (Ctrl+Click on windows) on each key point to view top matching sentences.\nThen use back link (Alt+LeftArrow on windows) to go back.')
    run.font.size = Pt(10)
    id_to_kids = create_dict_to_list([(e['data']['target'], e['data']['source']) for e in edges])

    logging.info('Creating key points hierarchy')

    id_to_paragraph1 = {}
    ids_order = []
    get_hierarchical_bullets(document, root_ids, id_to_kids, id_to_node, id_to_paragraph1, id_to_n_matches_subtree, ids_order=ids_order)

    logging.info('Creating key points matches tables')

    if n_matches is None:
        heading = document.add_heading(f'\n\nAll matches per key point', 1)
    else:
        heading = document.add_heading(f'\n\nTop {n_matches} matches per key point', 1)
    set_heading(heading)

    id_to_paragraph2 = {}
    for id in ids_order:
        n = id_to_node[id]
        kp = n["data"]["kp"]

        heading = document.add_heading(f'\nKey point: {kp}  ({n["data"]["n_matches"]} matches)', 2)
        set_heading(heading)
        id_to_paragraph2[id] = heading

        matches_dicts = kp_to_dicts[kp]
        if n_matches is not None and n_matches < len(kp_to_dicts[kp]):
            matches_dicts = sample_list_keep_order(matches_dicts, n_matches)

        logging.info(f'creating table for KP: {kp}, n_matches: {len(matches_dicts)}')

        records = []
        for d in matches_dicts:
            records.append([d["sentence_text"], trunc_float(float(d["match_score"]), 4)])

        heading = document.add_heading(f'Matching Sentences', 3)
        set_heading(heading)

        if include_match_score:
            table = document.add_table(rows=1, cols=2)
        else:
            table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        start = True
        for r in records:
            if start:
                row_cells = table.rows[0].cells
                start = False
            else:
                row_cells = table.add_row().cells
            row_cells[0].text = r[0]
            row_cells[0].width = Inches(5)
            if include_match_score:
                row_cells[1].text = str(r[1])
                row_cells[1].width = Inches(0.5)

    # add a bookmark to every paragraph
    for id, paragraph in id_to_paragraph2.items():
        add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=f'table_bookmark{id}')

    for id, paragraph in id_to_paragraph1.items():
        node = id_to_node[id]
        kp = node['data']['kp']
        n_matches = int(node["data"]["n_matches"])
        if n_matches == id_to_n_matches_subtree[id]:
            msg = f'{kp} ({n_matches} matches)'
        else:
            if sort_by_subtree:
                msg = f'{kp} ({id_to_n_matches_subtree[id]} matching sentences in subtree, {n_matches} matches)'
            else:
                msg = f'{kp} ({n_matches} matches, {id_to_n_matches_subtree[id]} in subtree)'
        add_link(paragraph=paragraph, link_to=f'table_bookmark{id}', text=msg, tool_tip="Click to view top matching sentences")

    for id, paragraph in id_to_paragraph1.items():
        add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=f'hierarchy_bookmark{id}')

    for id, paragraph in id_to_paragraph2.items():
        msg = ' - back'
        add_link(paragraph=paragraph, link_to=f'hierarchy_bookmark{id}', text=msg, tool_tip="Click to view hierarchy", set_color=True)

    logging.info(f'saving docx summary in file: {result_filename}')
    document.save(result_filename)
