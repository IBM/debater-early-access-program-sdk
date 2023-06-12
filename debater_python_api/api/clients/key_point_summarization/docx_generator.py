import logging
import random

import docx
import numpy as np
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches, Pt, RGBColor
from docx import Document

from debater_python_api.api.clients.key_point_summarization.KpsExceptions import KpsIllegalInputException
from debater_python_api.api.clients.key_point_summarization.utils import create_dict_to_list, \
    trunc_float, read_dicts_from_df, get_unique_sent_id

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

stance_to_stance_str = {"con": "Negative", "pro":"Positive", "":""}

def get_stance_to_stance_str(stances):
    if not stances:
        return ""
    if isinstance(stances, str):
        stances = [stances]
    if set(stances) == set(["no-stance"]):
        return ""
    if len(stances) == 1:
        return stance_to_stance_str[list(stances)[0]]
    if set(stances) == set(["pro","con"]):
        return "Positive and Negative"
    raise KpsIllegalInputException(f"unsupported stances {stances}")

def get_kp_stance_if_needed(id_data, stances):
    return id_data["kp_stance"].capitalize() + ", " if set(stances) == {"pro", "con"} else ""

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

def add_link(paragraph, link_to, text, tool_tip=None, set_color=False):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to,)

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip,)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.bold = False
    r.italic = False
    if set_color:
        r.font.color.rgb = RGBColor(0x40, 0x5E, 0x8D)

def set_heading(heading):
    paragraph_format = heading.paragraph_format
    paragraph_format.page_break_before = False
    paragraph_format.keep_with_next = False
    paragraph_format.keep_together = False

def add_data_stats(meta_data, dicts, kp_id_to_data, document, stances, min_n_matches):
    n_kps = len(kp_id_to_data)
    n_total_sentences = meta_data["general"]["n_sentences"]
    n_total_comments = meta_data["general"]["n_comments"]
    dicts_not_none = list(filter(lambda r: r["kp"] != "none", dicts))
    n_matches_comments = len(set([d["comment_id"] for d in dicts_not_none])) # Todo: this includes matches to smaller, filtered kps!
    rate_matched_comments = 100*n_matches_comments / n_total_comments

    heading = document.add_heading('Data Statistics', 1)
    set_heading(heading)
    s = f'Analyzed {n_total_comments} comments with {n_total_sentences} sentences.\n' \
        f'Identified {n_kps} key points.\n' \
        f'{int(np.round(rate_matched_comments, 0))}% of the comments were matched to at least one key point.\n'
    for stance in ["pro","con"]:
        if stance in stances:
            stance_str = get_stance_to_stance_str(stance).lower()
            n_kps_stance = len(list(filter(lambda kp_id:kp_id_to_data[kp_id].get("kp_stance")==stance, kp_id_to_data)))
            kp_stance_str = f"of {n_kps_stance} {stance_str} key points." if len(stances) > 1 else "key point."

            n_stance_comments = meta_data["per_stance"][stance]["n_comments_stance"]
            dicts_stance = list(filter(lambda r: r["kp_stance"] == stance, dicts_not_none))
            n_matches_comments_stance = len(set([d["comment_id"] for d in dicts_stance]))
            matched_comments_stance_rate = 100*n_matches_comments_stance / n_stance_comments if n_stance_comments > 0 else 0
            s += f'\nIdentified {n_stance_comments} {stance_str} comments.\n'\
                 f'{int(matched_comments_stance_rate)}% of these comments were matched to at least one {kp_stance_str}\n'
    p = document.add_paragraph()
    run = p.add_run(s)
    run.font.size = Pt(12)


def sample_list_keep_order(list_to_sample, n_to_sample, seed=0):
    rng = random.Random(seed)
    return [list_to_sample[i] for i in sorted(rng.sample(range(len(list_to_sample)), n_to_sample))]


def save_hierarchical_graph_data_to_docx(full_result_df, kp_id_to_data, result_filename, meta_data, n_matches=None, sort_by_subtree=True, include_match_score=False, min_n_matches=5, seed=0):

    def get_hierarchical_bullets_aux(document, kp_id_to_data, id, tab, id_to_paragraph, sort_by_subtree=True, ids_order=[]):
        bullet = '\u25E6' if tab % 2 == 1 else '\u2022'
        msg = f'{("   " * tab)} {bullet} '

        heading = document.add_heading(msg, 9 if (tab + 1) > 9 else (tab+1))
        set_heading(heading)

        id_to_paragraph[id] = heading
        ids_order.append(id)

        kids = kp_id_to_data[id].get('kids')
        if kids and len(kids) > 0:
            if sort_by_subtree:
                kids = sorted(kids, key=lambda n: int(kp_id_to_data[n].get('n_matching_sents_in_subtree')), reverse=True)
            else:
                kids = sorted(kids, key=lambda n: int(kp_id_to_data[n]['n_matching_sentences']), reverse=True)
            for k in kids:
                get_hierarchical_bullets_aux(document, kp_id_to_data, k, tab + 1, id_to_paragraph, sort_by_subtree, ids_order)


    def get_hierarchical_bullets(document, roots, kp_id_to_data, id_to_paragraph, sort_by_subtree=True, ids_order=[]):
        tab = 0
        if sort_by_subtree:
            roots = sorted(roots, key=lambda n: int(kp_id_to_data[n].get('n_matching_sents_in_subtree')), reverse=True)
        else:
            roots = sorted(roots, key=lambda n: int(kp_id_to_data[n]['n_matching_sentences']), reverse=True)
        for root in roots:
            get_hierarchical_bullets_aux(document, kp_id_to_data, root, tab, id_to_paragraph, sort_by_subtree=True, ids_order=ids_order)

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'


    heading = document.add_heading('Key Point Summarization Results', 1)
    set_heading(heading)

    if min_n_matches is not None:
        kp_id_to_data = {kp_id:kp_id_to_data[kp_id] for kp_id in kp_id_to_data if int(kp_id_to_data[kp_id]['n_matching_sentences'])>=min_n_matches}

    stances = list(meta_data["per_stance"].keys())#set(data['kp_stance'] for id,data in kp_id_to_data.items() if data['kp_stance'])
    stance_str = get_stance_to_stance_str(stances)
    if stance_str:
        p = document.add_paragraph(stance_str + " Key Points")
        p.style = document.styles['Subtitle']
        set_heading(p)
    insertHR(heading)

    kps = [data["kp"] for data in kp_id_to_data.values()]
    dicts, _ = read_dicts_from_df(full_result_df)
    dicts = list(filter(lambda d: d["kp"] in kps, dicts))
    kp_to_dicts = create_dict_to_list([(d['kp'], d) for d in dicts])

    add_data_stats(meta_data, dicts, kp_id_to_data, document, stances, min_n_matches)

    heading = document.add_heading('Key Point Hierarchy', 1)
    set_heading(heading)

    p = document.add_paragraph()
    run = p.add_run('Click (Ctrl+Click on windows) on each key point to view top matching sentences.\nThen use back link (Alt+LeftArrow on windows) to go back.')
    run.font.size = Pt(10)

    logging.info('Creating key points hierarchy')

    id_to_paragraph1 = {}
    ids_order = []

    root_ids = list(filter(lambda x:not kp_id_to_data[x].get('parent'), kp_id_to_data.keys()))
    get_hierarchical_bullets(document, root_ids, kp_id_to_data, id_to_paragraph1, ids_order=ids_order)

    logging.info('Creating key points matches tables')

    if n_matches is None:
        heading = document.add_heading(f'\n\nAll matches per key point', 1)
    else:
        heading = document.add_heading(f'\n\nTop {n_matches} matches per key point', 1)
    set_heading(heading)

    id_to_paragraph2 = {}
    for id in ids_order:
        id_data = kp_id_to_data[id]
        kp = id_data['kp']

        kp_stance_str = get_kp_stance_if_needed(id_data, stances)
        heading = document.add_heading(f'\nKey point: {kp}  ({kp_stance_str}{id_data["n_matching_comments"]}'
                                       f' matching comments, {id_data["n_matching_sentences"]} matching sentences)', 2)
        set_heading(heading)
        id_to_paragraph2[id] = heading

        matches_dicts = kp_to_dicts[kp]
        if n_matches is not None and n_matches < len(kp_to_dicts[kp]):
            matches_dicts = sample_list_keep_order(matches_dicts, n_matches, seed=seed)

        #logging.info(f'creating table for KP: {kp}, {stance_str}, matching sentences: {len(matches_dicts)}')

        records = []
        for d in matches_dicts:
            records.append([d["sentence_text"], trunc_float(float(d["match_score"]), 4)])

        heading = document.add_heading(f'Matching Sentences', 3)
        set_heading(heading)

        if include_match_score:
            table = document.add_table(rows=1, cols=2)
        else:
            table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'

        start = True
        for r in records:
            if start:
                row_cells = table.rows[0].cells
                start = False
            else:
                row_cells = table.add_row().cells
            row_cells[0].text = r[0]
            row_cells[0].width = Inches(5)
            if include_match_score:
                row_cells[1].text = str(r[1])
                row_cells[1].width = Inches(0.5)

    # add a bookmark to every paragraph
    for id, paragraph in id_to_paragraph2.items():
        add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=f'table_bookmark{id}')

    for id, paragraph in id_to_paragraph1.items():
        id_data = kp_id_to_data[id]
        kp = id_data['kp']
        kp_stance_str = get_kp_stance_if_needed(id_data, stances)
        n_matches = int(id_data["n_matching_sentences"])
        if n_matches == id_data["n_matching_sents_in_subtree"]:
            msg = f'{kp} ({kp_stance_str}{n_matches} matches)'
        else:
            if sort_by_subtree:
                msg = f'{kp} ({kp_stance_str}{id_data["n_matching_sents_in_subtree"]} matching sentences in subtree, {n_matches} matches)'
            else:
                msg = f'{kp} ({kp_stance_str}{n_matches} matches, {id_data["n_matching_sents_in_subtree"]} in subtree)'
        add_link(paragraph=paragraph, link_to=f'table_bookmark{id}', text=msg, tool_tip="Click to view top matching sentences")

    for id, paragraph in id_to_paragraph1.items():
        add_bookmark(paragraph=paragraph, bookmark_text="", bookmark_name=f'hierarchy_bookmark{id}')

    for id, paragraph in id_to_paragraph2.items():
        msg = ' - back'
        add_link(paragraph=paragraph, link_to=f'hierarchy_bookmark{id}', text=msg, tool_tip="Click to view hierarchy", set_color=True)

    logging.info(f'saving docx summary in file: {result_filename}')
    document.save(result_filename)
